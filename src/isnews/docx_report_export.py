"""Экспорт DOCX-отчета по результатам текущей сессии."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document

from src.isnews.config import PROJECT_PATHS, ProjectPaths


class DocxReportExportError(ValueError):
    """Ошибка экспорта DOCX-отчета."""


@dataclass(frozen=True)
class DocxReportExportResult:
    """Возвращает путь к DOCX-файлу и список включенных разделов."""

    report_path: Path
    generated_sections: tuple[str, ...]


def _get_available_path(target_path: Path) -> Path:
    """Подбирает свободный путь к файлу, если имя уже занято."""
    if not target_path.exists():
        return target_path

    counter = 1
    while True:
        candidate = target_path.with_name(
            f"{target_path.stem}_{counter}{target_path.suffix}"
        )
        if not candidate.exists():
            return candidate
        counter += 1


def _safe_text(value: Any) -> str:
    """Преобразует произвольное значение в строку."""
    if value is None:
        return "нет данных"
    return str(value)


def _add_key_value_rows(document: Document, rows: list[tuple[str, Any]]) -> None:
    """Добавляет в документ набор пар ключ-значение."""
    for key, value in rows:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{_safe_text(key)}: ").bold = True
        paragraph.add_run(_safe_text(value))


def _add_dataframe_table(
    document: Document,
    dataframe: pd.DataFrame,
    *,
    max_rows: int = 20,
) -> None:
    """Добавляет DataFrame как таблицу в DOCX-документ."""
    if dataframe.empty:
        document.add_paragraph("Данные отсутствуют.")
        return

    preview = dataframe.head(max_rows).copy()
    table = document.add_table(rows=1, cols=len(preview.columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for column_index, column_name in enumerate(preview.columns):
        header_cells[column_index].text = str(column_name)

    for row in preview.itertuples(index=False, name=None):
        cells = table.add_row().cells
        for cell_index, value in enumerate(row):
            cells[cell_index].text = _safe_text(value)


def _add_training_section(document: Document, training_result: Any) -> None:
    """Добавляет раздел по обучению модели."""
    document.add_heading("Обучение модели", level=1)
    if training_result is None:
        document.add_paragraph("Результат обучения в текущей сессии отсутствует.")
        return

    report = training_result.report
    model_name = getattr(report, "model_name", type(training_result.model).__name__)
    _add_key_value_rows(
        document,
        [
            ("Модель", model_name),
            ("Количество классов", len(getattr(report, "class_labels", ()))),
            ("Train строк", getattr(report, "train_rows", "")),
            ("Validation строк", getattr(report, "validation_rows", "")),
            ("Test строк", getattr(report, "test_rows", "")),
            ("Время обучения, сек", getattr(report, "training_seconds", "")),
            ("Файл модели", training_result.paths.model_path),
            ("JSON-отчет обучения", training_result.paths.report_path),
        ],
    )


def _add_evaluation_section(document: Document, evaluation_result: Any) -> None:
    """Добавляет раздел по метрикам качества."""
    document.add_heading("Метрики качества", level=1)
    if evaluation_result is None:
        document.add_paragraph("Метрики в текущей сессии отсутствуют.")
        return

    report = evaluation_result.report
    metrics_dataframe = pd.DataFrame(
        [
            {
                "Выборка": "train",
                "Accuracy": report.train_metrics.accuracy,
                "Precision macro": report.train_metrics.precision_macro,
                "Recall macro": report.train_metrics.recall_macro,
                "F1 macro": report.train_metrics.f1_macro,
            },
            {
                "Выборка": "validation",
                "Accuracy": report.validation_metrics.accuracy,
                "Precision macro": report.validation_metrics.precision_macro,
                "Recall macro": report.validation_metrics.recall_macro,
                "F1 macro": report.validation_metrics.f1_macro,
            },
            {
                "Выборка": "test",
                "Accuracy": report.test_metrics.accuracy,
                "Precision macro": report.test_metrics.precision_macro,
                "Recall macro": report.test_metrics.recall_macro,
                "F1 macro": report.test_metrics.f1_macro,
            },
        ]
    )
    document.add_paragraph(
        "В таблице ниже приведены основные метрики качества модели на обучающей, валидационной и тестовой выборках."
    )
    _add_dataframe_table(document, metrics_dataframe)
    document.add_paragraph(f"JSON-отчет: {_safe_text(evaluation_result.paths.report_path)}")
    for warning_message in report.warning_messages:
        document.add_paragraph(warning_message, style="List Bullet")


def _add_comparison_section(document: Document, comparison_result: Any) -> None:
    """Добавляет раздел по сравнению моделей."""
    document.add_heading("Сравнение моделей", level=1)
    if comparison_result is None:
        document.add_paragraph("Сравнение моделей в текущей сессии отсутствует.")
        return

    document.add_paragraph(
        f"Лучшая модель: {_safe_text(comparison_result.best_model_name or 'нет данных')}"
    )
    document.add_paragraph(f"CSV-отчет: {_safe_text(comparison_result.paths.csv_path)}")
    document.add_paragraph(f"JSON-отчет: {_safe_text(comparison_result.paths.json_path)}")
    _add_dataframe_table(document, comparison_result.dataframe)


def _add_registry_section(document: Document, registry_result: Any) -> None:
    """Добавляет раздел по реестру экспериментов."""
    document.add_heading("Реестр экспериментов", level=1)
    if registry_result is None:
        document.add_paragraph("Сводный реестр в текущей сессии отсутствует.")
        return

    document.add_paragraph(f"CSV-реестр: {_safe_text(registry_result.paths.csv_path)}")
    document.add_paragraph(f"JSON-реестр: {_safe_text(registry_result.paths.json_path)}")
    _add_dataframe_table(document, registry_result.dataframe)


def _add_error_analysis_section(document: Document, error_analysis_result: Any) -> None:
    """Добавляет раздел по анализу ошибок."""
    document.add_heading("Анализ ошибок", level=1)
    if error_analysis_result is None:
        document.add_paragraph("Анализ ошибок в текущей сессии отсутствует.")
        return

    report = error_analysis_result.report
    _add_key_value_rows(
        document,
        [
            ("Источник", report.source_name),
            ("Проверено строк", report.analyzed_rows),
            ("Ошибок", report.misclassified_rows),
            ("Корректных строк", report.correct_rows),
            ("Доля ошибок", report.error_rate),
            ("CSV ошибок", error_analysis_result.paths.misclassified_rows_path),
            ("JSON-отчет", error_analysis_result.paths.report_path),
        ],
    )
    _add_dataframe_table(document, error_analysis_result.misclassified_dataframe)


def export_session_docx_report(
    *,
    training_result: Any = None,
    evaluation_result: Any = None,
    comparison_result: Any = None,
    registry_result: Any = None,
    error_analysis_result: Any = None,
    project_paths: ProjectPaths = PROJECT_PATHS,
) -> DocxReportExportResult:
    """Экспортирует DOCX-отчет по результатам текущей сессии."""
    if all(
        item is None
        for item in (
            training_result,
            evaluation_result,
            comparison_result,
            registry_result,
            error_analysis_result,
        )
    ):
        raise DocxReportExportError(
            "Для DOCX-отчета пока нет данных. Сначала выполните хотя бы один этап обучения или анализа."
        )

    project_paths.ensure_directories()
    report_path = _get_available_path(project_paths.docx_reports_dir / "session_report.docx")

    document = Document()
    document.add_heading(
        "Отчет по интеллектуальному сервису классификации новостей",
        level=0,
    )
    document.add_paragraph(
        f"Сформировано: {datetime.now().astimezone().isoformat(timespec='seconds')}"
    )

    _add_training_section(document, training_result)
    _add_evaluation_section(document, evaluation_result)
    _add_comparison_section(document, comparison_result)
    _add_registry_section(document, registry_result)
    _add_error_analysis_section(document, error_analysis_result)

    document.save(report_path)

    generated_sections = tuple(
        section_name
        for section_name, result in (
            ("training", training_result),
            ("evaluation", evaluation_result),
            ("comparison", comparison_result),
            ("registry", registry_result),
            ("error_analysis", error_analysis_result),
        )
        if result is not None
    )
    return DocxReportExportResult(
        report_path=report_path,
        generated_sections=generated_sections,
    )
